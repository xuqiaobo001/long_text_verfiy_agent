#!/usr/bin/env python3
"""
提取合同审核要点并生成Excel格式的详细报告
"""

import sys
import re
from pathlib import Path
from collections import defaultdict, Counter
import pandas as pd

# 添加src目录到Python路径
sys.path.insert(0, str(Path(__file__).parent / 'src'))

# 尝试导入docx库
try:
    import docx
except ImportError:
    print("正在安装python-docx库...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

def load_docx_simple(file_path):
    """简单的DOCX文件读取"""
    doc = docx.Document(file_path)
    content = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            content.append(paragraph.text)

    # 处理表格
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                content.append(" | ".join(row_text))

    return '\n\n'.join(content)

def extract_key_contract_elements(text):
    """提取合同关键要素"""
    elements = {
        '合同主体': {},
        '金额条款': {},
        '时间节点': {},
        '质量标准': {},
        '违约条款': {},
        '付款条款': {},
        '验收条款': {},
        '变更条款': {},
        '争议解决': {},
        '生效条件': {},
        '其他重要条款': {}
    }

    # 1. 提取合同主体
    party_patterns = [
        r'(甲方|发包人|建设单位)[：:]?\s*([^，；。。\n]{1,100})',
        r'(乙方|承包人|施工单位)[：:]?\s*([^，；。。\n]{1,100})',
        r'(丙方|监理单位)[：:]?\s*([^，；。。\n]{1,100})',
        r'(联合体)[：:]?\s*([^，；。。\n]{1,200})'
    ]

    parties = []
    for pattern in party_patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            if len(match) >= 2:
                parties.append((match[0], match[1].strip()))

    elements['合同主体'] = {party_type: party_name for party_type, party_name in parties}

    # 2. 提取金额条款
    amount_patterns = [
        r'(合同总价|签约合同价|固定综合单价|单价|暂估价|暂列金额|保证金|违约金|赔偿金|罚金)[：:]?\s*([￥¥]?[\d,]+(?:\.\d{2})?(?:元|万元|亿元)?)',
        r'(人民币\s*[\d,]+(?:\.\d{2})?\s*元)',
        r'([￥¥]\s*[\d,]+(?:\.\d{2})?)'
    ]

    amounts = []
    for pattern in amount_patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            if isinstance(match, tuple) and len(match) >= 2:
                amounts.append((match[0].strip(), match[1].strip()))
            else:
                amounts.append(('金额', match.strip()))

    elements['金额条款'] = {amount_type: amount_value for amount_type, amount_value in amounts}

    # 3. 提取时间节点
    time_patterns = [
        r'(开工日期|计划开始工作日期|开工令日期)[：:]?\s*(\d{4}年\d{1,2}月\d{1,2}日|\d{4}-\d{1,2}-\d{1,2})',
        r'(竣工日期|计划竣工日期|交付日期)[：:]?\s*(\d{4}年\d{1,2}月\d{1,2}日|\d{4}-\d{1,2}-\d{1,2})',
        r'(工期|总工期)[：:]?\s*(\d+)(?:天|日|个月|年)',
        r'(缺陷责任期|保修期)[：:]?\s*(\d+)(?:天|日|个月|年)',
        r'(投标有效期|质保期)[：:]?\s*(\d+)(?:天|日|个月|年)',
        r'(付款期限|支付期限)[：:]?\s*(\d+)(?:天|日|个月|年)'
    ]

    times = []
    for pattern in time_patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            if len(match) >= 2:
                times.append((match[0].strip(), match[1].strip()))

    elements['时间节点'] = {time_type: time_value for time_type, time_value in times}

    # 4. 提取质量标准
    quality_patterns = [
        r'(质量标准|工程质量标准)[：:]?\s*([^，；。。\n]{1,200})',
        r'(合格|优良|不合格|符合.*标准|达到.*要求)',
        r'(设计质量标准)[：:]?\s*([^，；。。\n]{1,200})',
        r'(施工质量标准)[：:]?\s*([^，；。。\n]{1,200})'
    ]

    quality_standards = []
    for pattern in quality_patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            if isinstance(match, tuple) and len(match) >= 2:
                quality_standards.append((match[0].strip(), match[1].strip()))
            else:
                quality_standards.append(('质量标准', match.strip()))

    elements['质量标准'] = {q_type: q_value for q_type, q_value in quality_standards}

    # 5. 提取违约条款
    penalty_patterns = [
        r'(误期赔偿|逾期违约金|工期延误赔偿)[：:]?\s*([^，；。。\n]{1,200})',
        r'(违约金|赔偿金|罚金)[：:]?\s*([^，；。。\n]{1,200})',
        r'(.*违约责任)[：:]?\s*([^，；。。\n]{1,200})',
        r'(.*违约情形)[：:]?\s*([^，；。。\n]{1,200})'
    ]

    penalties = []
    for pattern in penalty_patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            if isinstance(match, tuple) and len(match) >= 2:
                penalties.append((match[0].strip(), match[1].strip()))

    elements['违约条款'] = {penalty_type: penalty_value for penalty_type, penalty_value in penalties}

    # 6. 提取付款条款
    payment_patterns = [
        r'(预付款|进度款|结算款|质保金)[：:]?\s*([^，；。。\n]{1,200})',
        r'(付款比例|支付比例)[：:]?\s*(\d+%)',
        r'(付款方式|支付方式)[：:]?\s*([^，；。。\n]{1,200})',
        r'(付款条件|支付条件)[：:]?\s*([^，；。。\n]{1,200})'
    ]

    payments = []
    for pattern in payment_patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            if isinstance(match, tuple) and len(match) >= 2:
                payments.append((match[0].strip(), match[1].strip()))

    elements['付款条款'] = {payment_type: payment_value for payment_type, payment_value in payments}

    # 7. 提取验收条款
    acceptance_patterns = [
        r'(竣工验收|单位验收|分部验收)[：:]?\s*([^，；。。\n]{1,200})',
        r'(验收标准|验收条件)[：:]?\s*([^，；。。\n]{1,200})',
        r'(验收程序)[：:]?\s*([^，；。。\n]{1,200})'
    ]

    acceptance = []
    for pattern in acceptance_patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            if isinstance(match, tuple) and len(match) >= 2:
                acceptance.append((match[0].strip(), match[1].strip()))

    elements['验收条款'] = {accept_type: accept_value for accept_type, accept_value in acceptance}

    # 8. 提取争议解决
    dispute_patterns = [
        r'(争议解决|纠纷解决)[：:]?\s*([^，；。。\n]{1,200})',
        r'(仲裁|诉讼|调解)[：:]?\s*([^，；。。\n]{1,200})',
        r'(管辖法院|仲裁机构)[：:]?\s*([^，；。。\n]{1,100})'
    ]

    disputes = []
    for pattern in dispute_patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            if isinstance(match, tuple) and len(match) >= 2:
                disputes.append((match[0].strip(), match[1].strip()))

    elements['争议解决'] = {dispute_type: dispute_value for dispute_type, dispute_value in disputes}

    # 9. 提取生效条件
    effective_patterns = [
        r'(合同生效|生效条件)[：:]?\s*([^，；。。\n]{1,200})',
        r'(签字盖章|签署)[：:]?\s*([^，；。。\n]{1,200})'
    ]

    effective = []
    for pattern in effective_patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            if isinstance(match, tuple) and len(match) >= 2:
                effective.append((match[0].strip(), match[1].strip()))

    elements['生效条件'] = {eff_type: eff_value for eff_type, eff_value in effective}

    return elements

def find_element_sections(text, elements):
    """查找各要素在哪些章节中出现"""
    sections = {}

    # 分割章节
    chapter_pattern = r'(第[一二三四五六七八九十\d]+[章节条款部分]|^\d+\.[^\n]*|^[一二三四五六七八九十]+、[^\n]*)'
    chapter_matches = list(re.finditer(chapter_pattern, text, re.MULTILINE))

    for element_category, element_dict in elements.items():
        sections[element_category] = {}

        for element_key, element_value in element_dict.items():
            locations = []

            for i, match in enumerate(chapter_matches):
                start_pos = match.start()
                end_pos = chapter_matches[i + 1].start() if i + 1 < len(chapter_matches) else len(text)
                chapter_text = text[start_pos:end_pos]

                # 检查要素是否在该章节中
                if element_key in chapter_text or element_value in chapter_text:
                    chapter_title = match.group().strip()
                    if len(chapter_title) > 50:  # 截断过长的标题
                        chapter_title = chapter_title[:50] + "..."
                    locations.append(chapter_title)

            sections[element_category][element_key] = locations

    return sections

def generate_excel_report(elements, sections, output_file):
    """生成Excel格式的审核报告"""

    # 准备数据
    report_data = []

    for category, element_dict in elements.items():
        for key, value in element_dict.items():
            # 获取出现的章节
            locations = sections.get(category, {}).get(key, [])
            location_str = "; ".join(locations) if locations else "未明确"

            row = {
                '要素类别': category,
                '具体条款': key,
                '条款内容': value,
                '出现章节': location_str,
                '审核状态': '需核对',
                '风险等级': determine_risk_level(category, key),
                '审核建议': generate_suggestion(category, key, value)
            }
            report_data.append(row)

    # 创建DataFrame
    df = pd.DataFrame(report_data)

    # 创建Excel writer对象
    with pd.ExcelWriter(output_file, engine='openpyxl') as writer:
        # 写入主报告
        df.to_excel(writer, sheet_name='合同审核要点', index=False)

        # 调整列宽
        worksheet = writer.sheets['合同审核要点']
        for column in worksheet.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            worksheet.column_dimensions[column_letter].width = adjusted_width

    return len(report_data)

def determine_risk_level(category, key):
    """确定风险等级"""
    high_risk_keys = ['合同总价', '违约金', '违约责任', '争议解决', '合同生效']
    medium_risk_keys = ['工期', '付款', '验收', '质量标准']

    if any(risk_key in key for risk_key in high_risk_keys):
        return '高'
    elif any(risk_key in key for risk_key in medium_risk_keys):
        return '中'
    else:
        return '低'

def generate_suggestion(category, key, value):
    """生成审核建议"""
    suggestions = {
        '合同主体': '核实主体名称的准确性和完整性',
        '金额条款': '核对金额计算和表述的一致性',
        '时间节点': '确认时间节点的合理性和可执行性',
        '质量标准': '检查质量标准是否符合国家规定',
        '违约条款': '评估违约条款的合理性和可操作性',
        '付款条款': '核实付款条款的明确性和可执行性',
        '验收条款': '确认验收标准和程序的完整性',
        '争议解决': '检查争议解决方式的合法性和有效性',
        '生效条件': '核实合同生效条件的完整性'
    }

    base_suggestion = suggestions.get(category, '建议仔细核对该条款')

    # 特殊情况建议
    if '不一致' in value or '不同' in value:
        return f"{base_suggestion}（注意：存在不一致表述）"
    elif '另行约定' in value or '另行商定' in value:
        return f"{base_suggestion}（注意：存在待约定条款）"
    elif len(value) < 10:
        return f"{base_suggestion}（注意：条款内容过于简略）"
    else:
        return base_suggestion

def main():
    """主函数"""
    contract_file = Path("/root/long_text_review/test_file/sj25-250501-doc.docx")
    output_file = "/root/long_text_review/test_file/contract_audit_detailed.xlsx"

    print("=== 合同详细审核要点提取工具 ===\n")

    try:
        # 安装必要的库
        try:
            import pandas as pd
        except ImportError:
            print("正在安装pandas和openpyxl...")
            import subprocess
            subprocess.check_call([sys.executable, "-m", "pip", "install", "pandas openpyxl"])
            import pandas as pd

        # 加载合同文件
        print(f"正在读取文件: {contract_file}")
        text = load_docx_simple(str(contract_file))
        print(f"文件大小: {len(text):,} 字符\n")

        # 提取关键要素
        print("正在提取合同关键要素...")
        elements = extract_key_contract_elements(text)

        print("正在查找要素分布...")
        sections = find_element_sections(text, elements)

        # 生成Excel报告
        print("正在生成Excel审核报告...")
        element_count = generate_excel_report(elements, sections, output_file)

        print(f"\n=== 提取完成 ===")
        print(f"共提取 {element_count} 个审核要点")
        print(f"详细报告已保存到: {output_file}")

        # 显示统计信息
        print("\n=== 要素统计 ===")
        for category, element_dict in elements.items():
            print(f"{category}: {len(element_dict)} 项")

    except Exception as e:
        print(f"错误: {e}")
        import traceback
        traceback.print_exc()

if __name__ == '__main__':
    main()