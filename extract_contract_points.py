#!/usr/bin/env python3
"""
提取合同审核要点
"""

import sys
import re
from pathlib import Path
from collections import Counter

# 添加src目录到Python路径
sys.path.insert(0, str(Path(__file__).parent / 'src'))

# 尝试导入docx库
try:
    import docx
except ImportError:
    print("正在安装python-docx库...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

def load_docx_simple(file_path):
    """简单的DOCX文件读取"""
    doc = docx.Document(file_path)
    content = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            content.append(paragraph.text)

    # 处理表格
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                content.append(" | ".join(row_text))

    return '\n\n'.join(content)

def extract_contract_sections(text):
    """提取合同的各个章节"""
    sections = {}

    # 常见的章节标题模式
    section_patterns = [
        r'第[一二三四五六七八九十\d]+[章节条款部分][：:\s]*([^\n\r]+)',
        r'(\d+\.[\d\.]*)\s*([^\n\r]+)',
        r'([一二三四五六七八九十]+)、([^\n\r]+)',
        r'^([A-Z]+)[\s:]+([^\n\r]+)',
        r'【([^】]+)】',
        r'（([^）]+)）'
    ]

    current_section = "前言"
    current_content = []
    lines = text.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # 检查是否是章节标题
        is_section = False
        for pattern in section_patterns:
            match = re.match(pattern, line)
            if match:
                # 保存之前的章节
                if current_content:
                    sections[current_section] = '\n'.join(current_content)

                # 开始新章节
                if len(match.groups()) >= 2:
                    current_section = f"{match.group(1)} {match.group(2)}"
                else:
                    current_section = match.group(1)
                current_content = []
                is_section = True
                break

        if not is_section:
            current_content.append(line)

    # 保存最后一个章节
    if current_content:
        sections[current_section] = '\n'.join(current_content)

    return sections

def extract_key_terms(text):
    """提取关键术语"""
    # 提取金额
    amounts = re.findall(r'[¥$]\s*\d+(?:,\d{3})*(?:\.\d{2})?|人民币\s*\d+(?:,\d{3})*(?:\.\d{2})?\s*元|\d+(?:,\d{3})*(?:\.\d{2})?\s*元', text)

    # 提取日期
    dates = re.findall(r'\d{4}年\d{1,2}月\d{1,2}日|\d{4}-\d{1,2}-\d{1,2}|\d{1,2}/\d{1,2}/\d{4}', text)

    # 提取百分比
    percentages = re.findall(r'\d+(?:\.\d+)?%', text)

    # 提取合同主体
    parties = re.findall(r'(甲方|乙方|丙方|丁方)[：:]?\s*([^\n\r]+?)(?=[，；。]|$)', text)

    # 提取法律法规
    laws = re.findall(r'《([^》]+(?:法|条例|规定|办法|细则))》', text)

    # 提取地点
    locations = re.findall(r'(北京|上海|广州|深圳|杭州|成都|武汉|西安|南京|天津|重庆|青岛|大连|厦门|苏州|无锡|常州|宁波|合肥|长沙|郑州|济南|福州|石家庄|昆明|南昌贵阳|南宁|太原、沈阳、长春、哈尔滨、呼和浩特、银川、西宁、乌鲁木齐、拉萨、兰州)[\u4e00-\u9fff]*', text)

    return {
        'amounts': amounts,
        'dates': dates,
        'percentages': percentages,
        'parties': parties,
        'laws': laws,
        'locations': locations
    }

def extract_review_points(text, sections, key_terms):
    """提取合同审核要点"""
    review_points = []

    # 1. 合同主体检查
    parties = key_terms['parties']
    if parties:
        party_names = [name[1].strip() for name in parties]
        unique_parties = len(set(party_names))
        if len(parties) != unique_parties:
            review_points.append({
                'type': '主体一致性',
                'priority': '严重',
                'issue': '合同主体名称存在不一致',
                'detail': f'发现 {len(parties)} 处提及，但只有 {unique_parties} 个不同的主体名称'
            })

    # 2. 必要条款检查
    required_clauses = ['违约', '争议解决', '生效', '终止', '保密', '知识产权']
    missing_clauses = []

    for clause in required_clauses:
        found = False
        for section_name, section_content in sections.items():
            if clause in section_name or clause in section_content:
                found = True
                break
        if not found:
            missing_clauses.append(clause)

    if missing_clauses:
        review_points.append({
            'type': '条款完整性',
            'priority': '重要',
            'issue': '缺少必要的合同条款',
            'detail': f'建议添加以下条款：{", ".join(missing_clauses)}'
        })

    # 3. 金额一致性检查
    amounts = key_terms['amounts']
    if len(amounts) > 1:
        # 提取数字金额进行对比
        numeric_amounts = []
        for amount in amounts:
            # 提取数字部分
            num_match = re.search(r'[\d,]+(?:\.\d+)?', amount.replace(',', ''))
            if num_match:
                numeric_amounts.append(float(num_match.group()))

        if len(set(numeric_amounts)) > 1:
            review_points.append({
                'type': '金额一致性',
                'priority': '重要',
                'issue': '合同中存在多个不同的金额',
                'detail': f'发现金额：{", ".join(amounts)}',
                'suggestion': '请核实各项金额的正确性'
            })

    # 4. 日期一致性检查
    dates = key_terms['dates']
    if len(dates) > 1:
        review_points.append({
            'type': '日期一致性',
            'priority': '中等',
            'issue': '合同包含多个日期',
            'detail': f'发现日期：{", ".join(dates[:5])}',
            'suggestion': '请确认所有日期的逻辑关系是否正确'
        })

    # 5. 模糊表述检查
    vague_patterns = [
        r'尽量|尽快|尽可能|相关|适当|合理|等',
        r'左右|大约|约|大概',
        r'有义务|有权|应承担|负责'
    ]

    vague_expressions = []
    for pattern in vague_patterns:
        matches = re.findall(pattern, text)
        vague_expressions.extend(matches)

    if vague_expressions:
        review_points.append({
            'type': '表述明确性',
            'priority': '中等',
            'issue': '存在模糊表述',
            'detail': f'发现模糊表述：{", ".join(list(set(vague_expressions))[:5])}',
            'suggestion': '建议将模糊表述替换为更具体的描述'
        })

    # 6. 数字一致性检查
    numbers = re.findall(r'\d+', text)
    number_counter = Counter(numbers)
    duplicate_numbers = {num: count for num, count in number_counter.items() if count > 3 and len(num) > 2}

    if duplicate_numbers:
        review_points.append({
            'type': '数字一致性',
            'priority': '低',
            'issue': '某些数字出现频率较高',
            'detail': f'请检查以下数字的一致性：{", ".join(list(duplicate_numbers.keys())[:5])}',
            'suggestion': '确保相同含义的数字在全文中保持一致'
        })

    # 7. 法律法规引用检查
    laws = key_terms['laws']
    if laws:
        review_points.append({
            'type': '法律依据',
            'priority': '信息',
            'issue': '引用了相关法律法规',
            'detail': f'引用法律：{", ".join(laws)}',
            'suggestion': '请确认引用的法律是否有效且适用'
        })

    # 8. 签字盖章检查
    if '签字' not in text and '盖章' not in text and '签署' not in text:
        review_points.append({
            'type': '签署条款',
            'priority': '重要',
            'issue': '未发现签字盖章相关条款',
            'suggestion': '应明确合同签署生效的条件'
        })

    return review_points

def main():
    """主函数"""
    contract_file = Path("/root/long_text_review/test_file/sj25-250501-doc.docx")

    print("=== 合同审核要点提取工具 ===\n")

    try:
        # 加载合同文件
        print(f"正在读取文件: {contract_file}")
        text = load_docx_simple(str(contract_file))

        print(f"文件大小: {len(text):,} 字符\n")

        # 提取章节
        print("正在分析合同结构...")
        sections = extract_contract_sections(text)
        print(f"发现 {len(sections)} 个章节:")
        for i, (title, content) in enumerate(sections.items(), 1):
            content_preview = content[:100].replace('\n', ' ') + "..." if len(content) > 100 else content
            print(f"  {i}. {title}")
            print(f"     {content_preview[:80]}...\n")

        # 提取关键信息
        print("\n=== 关键信息提取 ===")
        key_terms = extract_key_terms(text)

        if key_terms['parties']:
            print(f"\n合同主体:")
            for party_type, party_name in key_terms['parties']:
                print(f"  - {party_type}: {party_name}")

        if key_terms['amounts']:
            print(f"\n金额信息:")
            for amount in key_terms['amounts'][:10]:
                print(f"  - {amount}")

        if key_terms['dates']:
            print(f"\n日期信息:")
            for date in key_terms['dates'][:5]:
                print(f"  - {date}")

        if key_terms['laws']:
            print(f"\n法律法规引用:")
            for law in key_terms['laws']:
                print(f"  - {law}")

        # 提取审核要点
        print(f"\n=== 审核要点分析 ===\n")
        review_points = extract_review_points(text, sections, key_terms)

        if review_points:
            print(f"发现 {len(review_points)} 个审核要点:\n")

            # 按优先级分组
            priority_map = {'严重': 0, '重要': 1, '中等': 2, '低': 3, '信息': 4}
            review_points.sort(key=lambda x: priority_map.get(x['priority'], 3))

            for i, point in enumerate(review_points, 1):
                print(f"{i}. 【{point['priority']}】{point['type']}")
                print(f"   问题描述: {point['issue']}")
                if 'detail' in point:
                    print(f"   详细信息: {point['detail']}")
                if 'suggestion' in point:
                    print(f"   建议修改: {point['suggestion']}")
                print()
        else:
            print("未发现明显的审核问题。")

        # 保存提取的要点
        output_file = "/root/long_text_review/test_file/contract_review_points.txt"
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write("=== 合同审核要点报告 ===\n\n")

            f.write(f"文件: {contract_file}\n")
            f.write(f"文本长度: {len(text):,} 字符\n")
            f.write(f"章节数: {len(sections)}\n\n")

            f.write("=== 章节结构 ===\n")
            for title, content in sections.items():
                f.write(f"\n{title}\n")
                f.write(f"内容长度: {len(content)} 字符\n")

            f.write("\n=== 审核要点 ===\n\n")
            for i, point in enumerate(review_points, 1):
                f.write(f"{i}. 【{point['priority']}】{point['type']}\n")
                f.write(f"   问题描述: {point['issue']}\n")
                if 'detail' in point:
                    f.write(f"   详细信息: {point['detail']}\n")
                if 'suggestion' in point:
                    f.write(f"   建议修改: {point['suggestion']}\n")
                f.write("\n")

        print(f"\n审核要点已保存到: {output_file}")

    except Exception as e:
        print(f"错误: {e}")
        import traceback
        traceback.print_exc()

if __name__ == '__main__':
    main()